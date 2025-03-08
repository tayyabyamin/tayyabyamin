import random
import pyttsx3
import docx
import tkinter as tk
from tkinter import messagebox, ttk

# Initialize text-to-speech engine
engine = pyttsx3.init()
voices = engine.getProperty('voices')
rate = engine.getProperty('rate')
volume = engine.getProperty('volume')
current_voice = 0

# Function to read words from docx
def read_words_from_docx(filename):
    doc = docx.Document(filename)
    words = [para.text.strip().split()[0] for para in doc.paragraphs if para.text.strip() and not para.text.startswith("•")]
    return words

# Function to write words to docx
def write_words_to_docx(filename, words):
    doc = docx.Document()
    for word in words:
        doc.add_paragraph(word)
    doc.save(filename)

# Function to append words to a file
def append_to_file(filename, word):
    with open(filename, "a") as f:
        f.write(word + "\n")

# Function to speak a word
def speak(word):
    engine.say(word)
    engine.runAndWait()

# Function to get the next word
def next_word():
    global current_word
    if words:
        current_word = random.choice(words)
        speak(current_word)
        result_label.config(text="", fg="black")
        entry.delete(0, tk.END)
    else:
        result_label.config(text="No words left to test!", fg="black")

# Function to check spelling
def check_spelling():
    user_input = entry.get().strip()
    if not user_input:
        messagebox.showwarning("Warning", "Please enter a spelling!")
        return
    
    if user_input.lower() == current_word.lower():
        result_label.config(text=f"Correct spelling: {current_word}", fg="green")
        words.remove(current_word)
        write_words_to_docx("words.docx", words)
    else:
        incorrect_format = ''.join([f"\u0336{char}" for char in user_input])  # Apply strikethrough
        result_label.config(text=f"Wrong spelling: {incorrect_format}\nCorrect spelling: {current_word}", fg="red")
    
    append_to_file("correct_spellings.txt", current_word)
    speak(current_word)

# Function to repeat the current word
def repeat_word():
    speak(current_word)

# Function to update speech rate
def update_speed(val):
    global rate
    rate = int(float(val) * 100)
    engine.setProperty('rate', rate)
    speed_label.config(text=f"Speed: {float(val):.2f}x")

# Function to update volume
def update_volume(val):
    global volume
    volume = float(val)
    engine.setProperty('volume', volume)
    volume_label.config(text=f"Volume: {float(val):.2f}")

# Function to change the voice
def change_voice(event):
    selected_voice = voice_dropdown.get()
    for v in voices:
        if selected_voice in v.name:
            engine.setProperty('voice', v.id)
            break

# Load words
input_file = "words.docx"
words = read_words_from_docx(input_file)
current_word = ""

# GUI Setup
root = tk.Tk()
root.title("Spelling Test")
root.geometry("400x500")

instruction_label = tk.Label(root, text="Enter the spelling you hear:")
instruction_label.pack(pady=5)

entry = tk.Entry(root, font=("Arial", 14))
entry.pack(pady=5)

submit_button = tk.Button(root, text="Submit", command=check_spelling)
submit_button.pack(pady=5)

result_label = tk.Label(root, text="", font=("Arial", 12))
result_label.pack(pady=5)

repeat_button = tk.Button(root, text="Repeat Word", command=repeat_word)
repeat_button.pack(pady=5)

next_button = tk.Button(root, text="Next", command=next_word)
next_button.pack(pady=5)

speed_label = tk.Label(root, text=f"Speed: {rate/100:.2f}x")
speed_label.pack()
speed_slider = tk.Scale(root, from_=0.5, to=2.0, resolution=0.05, orient="horizontal", command=update_speed)
speed_slider.set(rate / 100)
speed_slider.pack()

volume_label = tk.Label(root, text=f"Volume: {volume:.2f}")
volume_label.pack()
volume_slider = tk.Scale(root, from_=0.0, to=1.0, resolution=0.05, orient="horizontal", command=update_volume)
volume_slider.set(volume)
volume_slider.pack()

voice_label = tk.Label(root, text="Select Voice:")
voice_label.pack()
voice_dropdown = ttk.Combobox(root, values=[v.name for v in voices])
voice_dropdown.pack()
voice_dropdown.bind("<<ComboboxSelected>>", change_voice)
voice_dropdown.current(0)  # Set default voice

next_word()
root.mainloop()