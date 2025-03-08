import random
import pyttsx3
import docx

def read_words_from_docx(filename):
    doc = docx.Document(filename)
    words = []
    for para in doc.paragraphs:
        line = para.text.strip()
        if line and not line.startswith("•"):  # Skip bullet points
            words.append(line.split()[0])  # Take only the first word
    return words

def write_words_to_docx(filename, words):
    doc = docx.Document()
    for word in words:
        doc.add_paragraph(word)
    doc.save(filename)

def append_to_file(filename, word):
    with open(filename, "a") as f:
        f.write(word + "\n")

def speak(word):
    engine = pyttsx3.init()
    engine.say(word)
    engine.runAndWait()

def spelling_test():
    input_file = "words.docx"
    correct_file = "correct_spellings.txt"
    
    words = read_words_from_docx(input_file)
    
    if not words:
        print("No words left to test!")
        return
    
    word = random.choice(words)
    
    while True:
        speak(word)
        user_input = input("Press Enter to reveal spelling or type 'r' to repeat: ").strip().lower()
        if user_input == 'r':
            continue
        else:
            break
    
    print("Correct spelling:", word)
    
    while True:
        user_input = input("Did you spell it correctly? (y/n): ").strip().lower()
        if user_input == 'y':
            append_to_file(correct_file, word)
            words.remove(word)
            write_words_to_docx(input_file, words)
            print("Word removed from list and added to correct spellings.")
            break
        elif user_input == 'n':
            print("Word remains in the list.")
            break
        else:
            print("Invalid input. Please enter 'y' or 'n'.")

if __name__ == "__main__":
    while True:
        spelling_test()
        if input("Press Enter to continue or type 'exit' to quit: ").strip().lower() == 'exit':
            break